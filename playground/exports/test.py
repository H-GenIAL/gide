from docx import Document
from docx.oxml.ns import qn
from utils import iterate_over_unique_cells, create_checkbox


def iterate_over_unique_cells_callback(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            # Find ##checkbox## in the text and replace it with a checkbox
            if "##checkbox##" in run.text:
                # run.text = "☐"
                run.text = "☑"

def main():
    doc = Document('./template.docx')
    for table in doc.tables:
        iterate_over_unique_cells(table, iterate_over_unique_cells_callback)
    doc.save("./output.docx")

if __name__ == "__main__":
    main()