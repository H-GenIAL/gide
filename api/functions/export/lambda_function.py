import boto3
import json
from docx import Document
from io import BytesIO
import base64
from pathlib import Path
import re

"""
Mapping dictionary that defines the structure and types of fields in the Word template.

Each key represents a field in the template that will be replaced with actual data.
The value is a dictionary containing the field type:
- "text": Simple text replacement
- "yesno": Yes/No checkbox replacement where one box will be checked based on the value
- "checkbox": Checkbox replacement where the checkbox will be checked if the value is "Oui"
- "date": Date field replacement
"""
cell_mapping = {
    "bailleur": {"type": "text"},
    "preneur": {"type": "text"},
    "cession": {"type": "yesno"},
    "adresse": {"type": "text"},
    "designation": {"type": "text"},
    "dest_loc_act": {"type": "text"},
    "clause_enseigne": {"type": "yesno"},
    "exclusivite": {"type": "yesno"},
    "non_concurrence": {"type": "yesno"},
    "classemt_loc_erp": {"type": "yesno"},
    "date_sign": {"type": "date"},
    "pinel": {"type": "yesno"},
    "date_prise_effet": {"type": "date"},
    "duree_bail": {"type": "text"},
    "terme_contrat_bail": {"type": "date"},
    "periode_ferme": {"type": "yesno"},
    "pro_fac_sortie": {"type": "date"},
    "preavis_min": {"type": "text"},
    "clause_spe_duree_bail_renouv": {"type": "yesno"},
    "mode_calc_fixe": {"type": "checkbox"},
    "mode_calc_paliers": {"type": "checkbox"},
    "mode_calc_recette": {"type": "checkbox"},
    "loyer_annuel_init": {"type": "text"},
    "loyer_annuel_cours": {"type": "text"},
    "paiement_trim_av": {"type": "yesno"},
    "tva": {"type": "yesno"},
    "clause_index": {"type": "yesno"},
    "date_index": {"type": "date"},
    "period_index": {"type": "text"},
    "indice_insee": {"type": "text"},
    "premiere_index": {"type": "text"},
    "index_suivantes": {"type": "text"},
    "indice_comparaison": {"type": "text"},
    "indice_base_fixe": {"type": "checkbox"},
    "index_hausse_uniqmt": {"type": "checkbox"},
    "plafond_plancher": {"type": "checkbox"},
    "risque_distorsion": {"type": "checkbox"},
    "divisibilite_clause_index": {"type": "checkbox"},
    "augment_comp_dern_loyer": {"type": "yesno"},
    "clause_loyer_bail_renouv": {"type": "yesno"},
    "franchise_reduct_loyer": {"type": "yesno"},
    "side_letter_tva": {"type": "text"},
    "autres_mesures_accomp": {"type": "yesno"},
    "depot_garantie": {"type": "yesno"},
    "autres_garanties": {"type": "yesno"},
    "garantie_possess_b": {"type": "yesno"},
    "garantie_cautio_soli": {"type": "checkbox"},
    "garantie_autonome": {"type": "checkbox"},
    "garantie_autre": {"type": "checkbox"},
    "garant_societe": {"type": "checkbox"},
    "garant_banque": {"type": "checkbox"},
    "montant": {"type": "text"},
    "expiration": {"type": "date"},
    "transfer_nb": {"type": "yesno"},
    "ppb_impots": {"type": "checkbox"},
    "ppb_taxe_fonc": {"type": "checkbox"},
    "ppb_teom": {"type": "checkbox"},
    "ppb_loc_comm": {"type": "checkbox"},
    "ppp_impots": {"type": "checkbox"},
    "ppp_taxe_fonc": {"type": "checkbox"},
    "ppp_teom": {"type": "checkbox"},
    "ppp_loc_comm": {"type": "checkbox"},
    "pp_impots": {"type": "checkbox"},
    "pp_taxe_fonc": {"type": "checkbox"},
    "pp_teom": {"type": "checkbox"},
    "pp_loc_comm": {"type": "checkbox"},
    "pcb_impots": {"type": "checkbox"},
    "pcb_taxe_fonc": {"type": "checkbox"},
    "pcb_teom": {"type": "checkbox"},
    "pcb_loc_comm": {"type": "checkbox"},
    "pcp_impots": {"type": "checkbox"},
    "pcp_taxe_fonc": {"type": "checkbox"},
    "pcp_teom": {"type": "checkbox"},
    "pcp_loc_comm": {"type": "checkbox"},
    "pc_impots": {"type": "checkbox"},
    "pc_taxe_fonc": {"type": "checkbox"},
    "pc_teom": {"type": "checkbox"},
    "pc_loc_comm": {"type": "checkbox"},
    "hb_gest_tech": {"type": "checkbox"},
    "hb_gest_loc": {"type": "checkbox"},
    "hb_gest_loyers": {"type": "checkbox"},
    "hb_gest_synd": {"type": "checkbox"},
    "hp_gest_tech": {"type": "checkbox"},
    "hp_gest_loc": {"type": "checkbox"},
    "hp_gest_loyers": {"type": "checkbox"},
    "hp_gest_synd": {"type": "checkbox"},
    "hnp_gest_tech": {"type": "checkbox"},
    "hnp_gest_loc": {"type": "checkbox"},
    "hnp_gest_loyers": {"type": "checkbox"},
    "hnp_gest_synd": {"type": "checkbox"},
    "assu_b_b": {"type": "checkbox"},
    "assu_b_p": {"type": "checkbox"},
    "assu_b_np": {"type": "checkbox"},
    "fonds_marktg": {"type": "yesno"},
    "pp_b_rep": {"type": "checkbox"},
    "pp_b_confo": {"type": "checkbox"},
    "pp_b_vetus": {"type": "checkbox"},
    "pp_b_equip": {"type": "checkbox"},
    "pp_p_rep": {"type": "checkbox"},
    "pp_p_confo": {"type": "checkbox"},
    "pp_p_vetus": {"type": "checkbox"},
    "pp_p_equip": {"type": "checkbox"},
    "pp_np_rep": {"type": "checkbox"},
    "pp_np_confo": {"type": "checkbox"},
    "pp_np_vetus": {"type": "checkbox"},
    "pp_np_equip": {"type": "checkbox"},
    "pc_b_rep": {"type": "checkbox"},
    "pc_b_confo": {"type": "checkbox"},
    "pc_b_vetus": {"type": "checkbox"},
    "pc_b_equip": {"type": "checkbox"},
    "pc_p_rep": {"type": "checkbox"},
    "pc_p_confo": {"type": "checkbox"},
    "pc_p_vetus": {"type": "checkbox"},
    "pc_p_equip": {"type": "checkbox"},
    "pc_np_rep": {"type": "checkbox"},
    "pc_np_confo": {"type": "checkbox"},
    "pc_np_vetus": {"type": "checkbox"},
    "pc_np_equip": {"type": "checkbox"},
    "derog_1722": {"type": "yesno"},
    "facu_modif": {"type": "yesno"},
    "facu_trav": {"type": "yesno"},
    "facu_p_trav": {"type": "yesno"},
    "facu_p_plaques": {"type": "yesno"},
    "etat_neuf": {"type": "checkbox"},
    "etat_parfait": {"type": "checkbox"},
    "etat_bon": {"type": "checkbox"},
    "etat_usage": {"type": "checkbox"},
    "clause_accession": {"type": "yesno"},
    "facu_b_remise_etat": {"type": "yesno"},
    "indemn_immo": {"type": "yesno"},
    "edl_entree_dataroom": {"type": "yesno"},
    "facu_sousloc": {"type": "yesno"},
    "facu_loca_gerance": {"type": "yesno"},
    "facu_domic": {"type": "yesno"},
    "facu_cess_droit_bail": {"type": "yesno"},
    "facu_cess_fonds_com": {"type": "yesno"},
    "locaux_indiv": {"type": "yesno"},
    "garanties_cession": {"type": "yesno"},
    "gar_cess_enti_prem": {"type": "checkbox"},
    "gar_cess_enti_successifs": {"type": "checkbox"},
    "gar_cess_dur_stip_3y": {"type": "yesno"},
    "gar_cess_cessionnaire": {"type": "yesno"},
    "droit_pref_b_app": {"type": "yesno"},
    "droit_pref_b_loc_loues": {"type": "yesno"},
    "droit_pref_b_immeuble": {"type": "yesno"},
    "droit_pref_p_fonds_com": {"type": "yesno"},
    "annex_dta": {"type": "yesno"},
    "annex_er": {"type": "yesno"},
    "annex_dpe": {"type": "yesno"},
    "annex_envir": {"type": "yesno"},
    "decret_terti_applicable": {"type": "yesno"},
    "icpe": {"type": "yesno"},
    "stipu_intuitu_personae": {"type": "yesno"},
    "renonc_imprev": {"type": "yesno"},
    "relations_impayes": {"type": "yesno"},
    "relations_ech_signif": {"type": "yesno"},
    "relations_precontentieux": {"type": "yesno"},
    "commentaires": {"type": "text"},
    "index_audites": {"type": "text"}
}

def docx_iterate_over_unique_cells(table, func):
    """
    Iterate over all unique cells in a table (ignoring merged/extended cells)
    and apply a function to each cell.
    """
    # Get unique cells (ignoring merged/extended cells)
    cells = set(cell for row in table.rows for cell in row.cells)
    for cell in cells:
        func(cell)

def docx_replace_text(paragraphs, text, value):
    """
    Replace a text found inside a paragraph with a value.
    """
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if text in run.text:
                run.text = value

def replace_pattern_in_paragraph(paragraph, data):
    """
    Replace patterns in a paragraph with their corresponding values.
    Handles both simple text replacements and yes/no checkboxes.
    """
    for run in paragraph.runs:
        # Find all patterns like $key$ or $key-yes$ $key-no$
        patterns = re.finditer(r'\$([\w-]+)\$', run.text)
        for match in patterns:
            full_pattern = match.group(0)  # The full pattern including $
            key = match.group(1)  # The key without $
            
            # Handle yes/no patterns
            if key.endswith('-oui') or key.endswith('-non'):
                base_key = key[:-4]  # Remove -oui or -non
                is_yes = key.endswith('-oui')
                if base_key in data:
                    value = data[base_key]
                    # Check for both English and French positive values
                    is_positive = value.lower() in ['oui', 'yes', 'true']
                    run.text = run.text.replace(full_pattern, "☑" if is_positive == is_yes else "☐")
            # Handle checkbox patterns
            elif key in cell_mapping and cell_mapping[key]['type'] == 'checkbox':
                value = data.get(key, '')
                run.text = run.text.replace(full_pattern, "☑" if value == "Oui" else "☐")
            # Handle regular text patterns
            elif key in data:
                run.text = run.text.replace(full_pattern, str(data[key]))

def export_to_word_base64(data, template_path):
    """
    Export the structured data to a Word document and return the base64 encoded document.

    Args:
        data (dict): The structured data to export.
        template_path (str): The path to the Word template.

    Returns:
        str: The base64 encoded document.
    """
    doc = Document(template_path)
    
    # Process all paragraphs in the document
    for paragraph in doc.paragraphs:
        replace_pattern_in_paragraph(paragraph, data)
    
    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_pattern_in_paragraph(paragraph, data)
    
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Convert the document to base64
    base64_doc = base64.b64encode(doc_buffer.getvalue()).decode('utf-8')
    return base64_doc

def lambda_handler(event, context):
    """
    Lambda function to export a Word document from a structured data.

    Args:
        event (dict): The event object containing the request data.
        context (dict): The context object containing the runtime information.
    """
    s3_client = boto3.client('s3')
    
    try:
        # Get the bucket name and object key
        bucket_name = 'apigide-exports'
        object_key = 'template.docx'
        
        # Get the object from S3
        response = s3_client.get_object(
            Bucket=bucket_name,
            Key=object_key
        )
        
        # Read the content
        template_doc_content = response['Body'].read()
        template_doc_buffer = BytesIO(template_doc_content)

        request_data = json.loads(event['body']) if isinstance(event.get('body'), str) else event.get('body', {})

        output_doc_base64 = export_to_word_base64(request_data, template_doc_buffer)
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'Content-Disposition': 'attachment; filename="document.docx"',
                'Access-Control-Allow-Origin': '*'
            },
            'body': output_doc_base64,
            'isBase64Encoded': True
        }
    
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({
                'message': f'Error: {str(e)}'
            })
        }

def save_to_file(base64_content, output_path):
    """
    Save a base64 encoded document to a file.
    """
    # Create the subdirectory if it doesn't exist of the output path
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    # Decode the base64 string back to bytes before writing
    decoded_content = base64.b64decode(base64_content)
    with open(output_path, 'wb') as file:
        file.write(decoded_content)

def debug_lambda_handler(event, context):
    """
    Debug the lambda function by exporting a Word document from a mocked structured data.

    Args:
        event (dict): The event object containing the request data.
        context (dict): The context object containing the runtime information.
    """
    try:
        template_doc_content = Path('data/template.docx').read_bytes()
        template_doc_buffer = BytesIO(template_doc_content)

        request_data = json.loads(event['body']) if isinstance(event.get('body'), str) else event.get('body', {})

        output_doc_base64 = export_to_word_base64(request_data, template_doc_buffer)

        save_to_file(output_doc_base64, "dist/output.docx")

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    with open("./data/input.json", "r") as file:
        input_data = json.load(file)

    debug_lambda_handler({
        "body": input_data
    }, None)