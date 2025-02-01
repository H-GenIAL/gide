import boto3
import json
from docx import Document
import sys
from io import BytesIO
import base64
from pathlib import Path

"""
Mapping the types of the cells in the template to the data in the YAML file
to be used in the yaml-to-word.py script to replace the template keys with the data
"""
cell_mapping = {
    "bailleur": {
        "type": "text"
    },
    "preneur": {
        "type": "text"
    },
    "cession": {
        "type": "yesno"
    },
    "cession_raison": {
        "type": "text"
    },
    "adresse": {
        "type": "text"
    },
    "designation": {
        "type": "yesno"
    }
}

def docx_iterate_over_unique_cells(table, func):
    """
    Iterate over all unique cells in a table (ignoring merged/extended cells)
    """
    # Get unique cells (ignoring merged/extended cells)
    cells = set(cell for row in table.rows for cell in row.cells)
    for cell in cells:
        func(cell)

def load_json(file_path):
    """Load and parse JSON file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return json.load(file)
    except Exception as e:
        print(f"Error loading JSON file: {e}")
        sys.exit(1)

def docx_replace_text(cell, text, value):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if text in run.text:
                run.text = value

def replace_key_with_value(key, value):
    """Create a callback function to replace template keys with values in Word doc cells."""
    type = cell_mapping[key]['type'] if key in cell_mapping else "text"
    
    def callback(cell):
        if type == "yesno":
            docx_replace_text(cell, f"${key}-yes", "☑" if value == "Yes" else "☐")
            docx_replace_text(cell, f"${key}-no", "☐" if value == "Yes" else "☑")
        else:
            docx_replace_text(cell, f"${key}", value)
    return callback

def export_to_word_base64(data, template_path):
    doc = Document(template_path)
    table = doc.tables[0]
    for key, value in data.items():
        # Iterate over all unique cells in the table and replace the key with the value
        docx_iterate_over_unique_cells(table, replace_key_with_value(key, value))
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Convert the document to base64
    base64_doc = base64.b64encode(doc_buffer.getvalue()).decode('utf-8')
    return base64_doc

def lambda_handler(event, context):
    s3_client = boto3.client('s3')
    
    try:
        # Get the bucket name and object key
        bucket_name = 'apigide-exports'
        object_key = 'template.docx'
        
        # Get the object from S3
        response = s3_client.get_object(
            Bucket=bucket_name,
            Key=object_key
        )
        
        # Read the content
        template_doc_content = response['Body'].read()
        template_doc_buffer = BytesIO(template_doc_content)

        request_data = json.loads(event['body']) if isinstance(event.get('body'), str) else event.get('body', {})

        output_doc_base64 = export_to_word_base64(request_data, template_doc_buffer)
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'Content-Disposition': 'attachment; filename="document.docx"',
                'Access-Control-Allow-Origin': '*'
            },
            'body': output_doc_base64,
            'isBase64Encoded': True
        }
    
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({
                'message': f'Error: {str(e)}'
            })
        }

def debug_lambda_handler(event, context):
    # Read the content
    try:
        template_doc_content = Path('data/template.docx').read_bytes()
        template_doc_buffer = BytesIO(template_doc_content)

        request_data = json.loads(event['body']) if isinstance(event.get('body'), str) else event.get('body', {})

        output_doc_base64 = export_to_word_base64(request_data, template_doc_buffer)
        
        print(output_doc_base64)
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    debug_lambda_handler({
        "body": {
            "bailleur": "John Doe",
            "preneur": "Jane Doe",
            "cession": "Yes",
            "cession_reason": "Test",
            "adresse": "123 Main St",
            "designation": "Yes"
        }
    }, None)