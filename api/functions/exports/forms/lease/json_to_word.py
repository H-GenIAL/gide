import json
from docx import Document
import argparse
import sys
from pathlib import Path
from mapping import cell_mapping
from utils import iterate_over_unique_cells
from io import BytesIO
import base64

def load_json(file_path):
    """Load and parse JSON file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return json.load(file)
    except Exception as e:
        print(f"Error loading JSON file: {e}")
        sys.exit(1)

def replace_key_with_value(key, value):
    """Create a callback function to replace template keys with values in Word doc cells."""
    if key not in cell_mapping:
        return None
    
    type = cell_mapping[key]['type']
    
    if type == "yesno":
        def callback(cell):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if f"${key}-yes" in run.text:
                        run.text = "☑" if value == "Yes" else "☐"
                    elif f"${key}-no" in run.text:
                        run.text = "☐" if value == "Yes" else "☑"
        return callback

    # text, number, etc.
    else:
        def callback(cell):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if f"${key}" in run.text:
                        run.text = value
        return callback

def export_to_word(data, output_path, template_path):
    """Write data to a Word document."""
    template_doc = Document(template_path)
    table = template_doc.tables[0]
    for key, value in data.items():
        # Iterate over all unique cells in the table and replace the key with the value
        iterate_over_unique_cells(table, replace_key_with_value(key, value))
    template_doc.save(output_path)

def export_to_word_base64(data, template_path):
    doc = Document(template_path)
    table = doc.tables[0]
    for key, value in data.items():
        # Iterate over all unique cells in the table and replace the key with the value
        iterate_over_unique_cells(table, replace_key_with_value(key, value))
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Convert the document to base64
    base64_doc = base64.b64encode(doc_buffer.getvalue()).decode('utf-8')
    return base64_doc

def main():
    parser = argparse.ArgumentParser(description='Convert JSON to Word document')
    parser.add_argument('-i', '--input', help='Input JSON file path')
    parser.add_argument('-t', '--template', help='Template Word document path')
    parser.add_argument('-o', '--output', help='Output Word document path')
    
    args = parser.parse_args()
    
    # Validate input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Input JSON file not found: {input_path}")
        sys.exit(1)

    template_path = Path(args.template or "./template.docx")
    if not template_path.exists():
        print(f"Template Word document not found: {template_path}")
        sys.exit(1)
    
    # Create output directory if it doesn't exist
    output_path = Path(args.output)
    if not output_path.exists():
        output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Process the conversion
    data = load_json(input_path)
    export_to_word(data, output_path, template_path)

if __name__ == '__main__':
    main()